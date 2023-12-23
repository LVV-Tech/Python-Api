from docx import Document
import aspose.words as aw

# Открываем документ
doc = Document("Согалсие_на_обработку_персональных_данных.docx")

# Получаем все параграфы документа
paras = doc.paragraphs

# Проходим по всем параграфам и заменяем необходимые поля
for para in paras:
    if "{{name}}" in para.text:
        para.text = para.text.replace("{{name}}", "Королев Максим Алексеевич")

# Сохраняем изменения
doc.save("output.docx")


doc = aw.Document("output.docx")
doc.save("out.pdf")


# convert("output.docx")
# convert("Согалсие_на_обработку_персональных_данных.docx", "output.pdf")
