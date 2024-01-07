import math
import telebot
import os
import storage
import re
from docx import Document
from datetime import datetime

token = "6417722044:AAFil8RguC7k2YUn-eyKG59yar-oSjgkpso"
bot = telebot.TeleBot(token)

services = ["дистрибьюция", "сведение и мастеринг", "создание дизайна", "аренда бита", "создание трека"]
servicesDescriptions = ["Дистрибуция - публикация Вашего трека/Альбома на все цифровые площадки.\nКаждый релиз отправляется на модерацию ПРОМО всех стриминговых сервисов.\nСсылка на портфолио - (позже)\n\n⏳ Сроки: от 3 до 5 дней.",
                        "Сведение и Мастеринг - выполнение комплекса работ по постобработке музыкального произведения.\nОчистка и подготовка звуковых дорожек, Обработка вокала, Питч-коррекция. Добавление различных эффектов и особенностей, Расширение стереобазы, работа с панорамой, Эквализация и компрессия,  и многое другое.\nВсе Ваши пожелания и требования будут учтены в процессе работы. Вам доступно Три бесплатных правки итоговой работы.\nСсылка на портфолио - (позже)\n\n⏳ Сроки: от 3 до 5 дней.",
                        "Создание дизайна - Создание обложки для Релиза/Альбома, Ретушь и обработка фото, создание дизайна для ПРОМО.\nВсе Ваши пожелания и требования будут учтены в процессе работы. Вам доступно Три бесплатных правки итоговой работы.\nСсылка на портфолио - https://vk.com/album-221967488_295470314\n\n⏳ Сроки: от 3 до 5 дней.",
                        "Аренда бита - любая категория аренды бита, от MP3 до EXCLUSIVE. Все биты представленные в плейлистах доступны для покупки.\nПлейлисты разбиты по разным Авторам и  жанрам, выбирайте подходящий Вам звук.\n\nСсылка на плейлист 1 - (позже)\n\n⏳ Сроки: 1 день.",
                        "Создание трека - услуга в полном объеме временно недоступна.\n\n⏳ Сроки: по договорённости"]
servicesPrices = [500, 2500, 1500, 1500, 5000]

prevService = 0

@bot.message_handler(commands=['start'])
def welcome(message):
    chat_id = message.chat.id
    keyboard = telebot.types.ReplyKeyboardMarkup()
    btn1 = telebot.types.KeyboardButton("Связаться с менеджером")
    keyboard.row(btn1)
    btn2 = telebot.types.KeyboardButton("Услуги")
    btn3 = telebot.types.KeyboardButton("Отзывы")
    keyboard.row(btn2, btn3)
    btn4 = telebot.types.KeyboardButton("О нас")
    btn5 = telebot.types.KeyboardButton("FAQ")
    keyboard.row(btn4, btn5)
    bot.send_message(chat_id,
                     'Вот что мы можем вам предложить',
                     reply_markup=keyboard)
    bot.register_next_step_handler(message, on_click)

@bot.message_handler(content_types=['text'])
def on_click(msg):
    global prevService
    chat_id = msg.chat.id
    msg.text = msg.text.lower()
    if msg.text == "меню":
        keyboard = telebot.types.ReplyKeyboardMarkup()
        btn1 = telebot.types.KeyboardButton("Связаться с менеджером")
        keyboard.row(btn1)
        btn2 = telebot.types.KeyboardButton("Услуги")
        btn3 = telebot.types.KeyboardButton("Отзывы")
        keyboard.row(btn2, btn3)
        btn4 = telebot.types.KeyboardButton("О нас")
        btn5 = telebot.types.KeyboardButton("FAQ")
        keyboard.row(btn4, btn5)
        bot.send_message(chat_id,
                        'Вот что мы можем вам предложить',
                        reply_markup=keyboard)
    if msg.text == "связаться с менеджером":
        keyboard = telebot.types.ReplyKeyboardMarkup()
        btn1 = telebot.types.KeyboardButton("Творческий/Примеры работ, оценка, услуги")
        keyboard.row(btn1)
        btn2 = telebot.types.KeyboardButton("Юридический вопрос/Оформление документов")
        keyboard.row(btn2)
        btn3 = telebot.types.KeyboardButton("Маркетинг/Реклама, продвижение")
        keyboard.row(btn3)
        btn4 = telebot.types.KeyboardButton("Технический вопрос/Не работает сайт, бот")
        keyboard.row(btn4)
        btn5 = telebot.types.KeyboardButton("Меню")
        keyboard.row(btn5)
        bot.send_message(chat_id,
                     'Мы прочитаем ваш вопрос и ответим на него как можно скорее. Для удобства выберите тему вопроса.',
                     reply_markup=keyboard)
    if msg.text == "творческий/примеры работ, оценка, услуги" or \
       msg.text == "юридический вопрос/оформление документов" or \
       msg.text == "маркетинг/реклама, продвижение" or \
       msg.text == "технический вопрос/не работает сайт, бот":
        keyboard = telebot.types.InlineKeyboardMarkup()
        keyboard.add(telebot.types.InlineKeyboardButton("Меню", callback_data='menu'))
        bot.send_message(chat_id,
                     'Чем я могу вам помочь? Опишите свою проблему следующим сообщением',
                     reply_markup=keyboard)
    if msg.text == "услуги" or msg.text == "⏪назад⏪":
        keyboard = telebot.types.ReplyKeyboardMarkup()

        btn1 = telebot.types.KeyboardButton("Дистрибьюция")
        keyboard.row(btn1)
        btn2 = telebot.types.KeyboardButton("Сведение и Мастеринг")
        keyboard.row(btn2)
        btn3 = telebot.types.KeyboardButton("Создание Дизайна")
        keyboard.row(btn3)
        btn4 = telebot.types.KeyboardButton("Аренда Бита")
        keyboard.row(btn4)
        btn5 = telebot.types.KeyboardButton("Создание трека")
        keyboard.row(btn5)
        btn6 = telebot.types.KeyboardButton("Узнать про пакеты")
        keyboard.row(btn6)
        btn7 = telebot.types.KeyboardButton("Меню")
        keyboard.row(btn7)
        bot.send_message(chat_id,
                     'Выберите услугу',
                     reply_markup=keyboard)
    for i in range(len(services)):
        if msg.text == services[i]:
            prevService = i
            keyboard = telebot.types.ReplyKeyboardMarkup()
            btn1 = telebot.types.KeyboardButton("✅Подтвердить заказ✅")
            keyboard.row(btn1)
            btn2 = telebot.types.KeyboardButton("⏪Назад⏪")
            keyboard.row(btn2)
            bot.send_message(chat_id,
                     f"{servicesDescriptions[i]}\n💸 Цена: от {servicesPrices[i]} ₽.",
                     reply_markup=keyboard)
            break
    if msg.text == "узнать про пакеты":
        keyboard = telebot.types.ReplyKeyboardMarkup()
        keyboard.row(telebot.types.KeyboardButton("Меню"))
        bot.send_message(chat_id,
                     'Ваш Запрос в работе, в ближайшее время с вами свяжется менеджер',
                     reply_markup=keyboard)
    if msg.text == "✅подтвердить заказ✅":
        keyboard = telebot.types.ReplyKeyboardMarkup()
        keyboard.row(telebot.types.KeyboardButton("Меню"))
        print("Вы выбрали услугу", services[prevService])
        if storage.check_user(tg_id=chat_id):
            bot.send_message(chat_id,
                     'Заполните документы и отправьте их модератору: @ghostikgh',
                     reply_markup=keyboard)
            # Открываем документ
            doc = Document("sogl.docx")
            user = storage.get_user_tg_id(chat_id)
            print(user)
            # Получаем все параграфы документа
            paras = doc.paragraphs
            
            today = datetime.now()
            

            # Проходим по всем параграфам и заменяем необходимые поля
            for para in paras:
                if "{{name}}" in para.text:
                    para.text = para.text.replace("{{name}}", user[0][2].title())
            for para in paras:
                if "{{passport}}" in para.text:
                    para.text = para.text.replace("{{passport}}", user[0][3])
            for para in paras:
                if "{{day}}" in para.text:
                    para.text = para.text.replace("{{day}}", datetime.strftime(today, '%d'))
            for para in paras:
                if "{{month}}" in para.text:
                    para.text = para.text.replace("{{month}}", datetime.strftime(today, ' %B'))
            for para in paras:
                if "{{year}}" in para.text:
                    para.text = para.text.replace("{{year}}", datetime.strftime(today, '%Y'))

            # Сохраняем изменения
            doc.save("output.docx")
            with open("output.docx", 'rb') as document:
                bot.send_document(chat_id, document)
            
            doc = Document("dog.docx")
            user = storage.get_user_tg_id(chat_id)
            # Получаем все параграфы документа
            paras = doc.paragraphs
            
            today = datetime.now()
            

            # Проходим по всем параграфам и заменяем необходимые поля
            for para in paras:
                if "{{name}}" in para.text:
                    para.text = para.text.replace("{{name}}", user[0][2].title())
            for para in paras:
                if "{{passport}}" in para.text:
                    para.text = para.text.replace("{{passport}}", user[0][3])
            for para in paras:
                if "{{day}}" in para.text:
                    para.text = para.text.replace("{{day}}", datetime.strftime(today, '%d'))
            for para in paras:
                if "{{month}}" in para.text:
                    para.text = para.text.replace("{{month}}", datetime.strftime(today, ' %B'))
            for para in paras:
                if "{{year}}" in para.text:
                    para.text = para.text.replace("{{year}}", datetime.strftime(today, '%Y'))
            for para in paras:
                if "{{services}}" in para.text:
                    para.text = para.text.replace("{{services}}", services[prevService])
            for para in paras:
                if "{{price}}" in para.text:
                    para.text = para.text.replace("{{price}}", str(servicesPrices[prevService]))

            # Сохраняем изменения
            doc.save("output.docx")
            with open("output.docx", 'rb') as document:
                bot.send_document(chat_id, document)
        else:
            keyboard.row(telebot.types.KeyboardButton("Регистрация"))
            bot.send_message(chat_id,
                     'Вы не зарегистрированы',
                     reply_markup=keyboard)
    if msg.text == "регистрация":
        keyboard = telebot.types.ReplyKeyboardMarkup()
        keyboard.row(telebot.types.KeyboardButton("Меню"))
        bot.send_message(chat_id, 'Введите свои паспортные данные в формате: серия номер', reply_markup=keyboard)

        @bot.message_handler(func=lambda message: re.match(r'^[0-9]{4} [0-9]{6}$', message.text))
        def passport(message):
            match = re.match(r'^[0-9]{4} [0-9]{6}$', message.text)
            if match:
                passp = message.text
                bot.send_message(chat_id, 'Введите свой номер телефона в формате: +79999999999', reply_markup=keyboard)
            else:
                bot.send_message(chat_id, 'Ошибка, проверьте правильность введенных данных', reply_markup=keyboard)
                return
            @bot.message_handler(func=lambda message: re.match(r'^\+7[0-9]{10}$', message.text))
            def phone(message):
                if match:
                    phone = message.text
                    bot.send_message(chat_id, 'Введите свою фамилию имя отчество:', reply_markup=keyboard)
                else:
                    bot.send_message(chat_id, 'Ошибка, проверьте правильность введенных данных', reply_markup=keyboard)
                    return

                @bot.message_handler(func=lambda message: re.match(r'^[^\W\d_]+\s[^\W\d_]+?(\s[^\W\d_]+)?$', message.text))
                def name(message):
                    
                    if match:
                        name = message.text
                        storage.create_user(
                            role=0,
                            phone=phone,
                            full_name=name,
                            passport=passp,
                            tg_id=chat_id
                        )
                        bot.send_message(chat_id, 'Вы зарегистрированы', reply_markup=keyboard)
                    else:
                        bot.send_message(chat_id, 'Ошибка, проверьте правильность введенных данных', reply_markup=keyboard)
                        return
                bot.register_next_step_handler(message, name)
                    
            @bot.message_handler(func=lambda message: not re.match(r'^\+7[0-9]{10}$', message.text))
            def invalid_phone(message):
                bot.send_message(chat_id, 'Ошибка, проверьте правильность введенных данных', reply_markup=keyboard)
                bot.register_next_step_handler(message, phone)

            bot.register_next_step_handler(message, phone)

        @bot.message_handler(func=lambda message: not re.match(r'^[0-9]{4} [0-9]{6}$', message.text))
        def invalid_passport(message):
            bot.send_message(chat_id, 'Ошибка, проверьте правильность введенных данных', reply_markup=keyboard)
            bot.register_next_step_handler(message, passport)

        bot.register_next_step_handler(msg, passport)
    if msg.text == "отзывы":
        bot.send_message(chat_id,
                     'Отзывы о нас вы можете найти [здесь](https://vk.com/topic-221967488_49474380)',
                     parse_mode='Markdown')
    if msg.text == "о нас":
        bot.send_message(chat_id,
                     'Прочитать о нас вы можете [здесь](https://vk.com/@lvvlabel-kto-my)',
                     parse_mode='Markdown') 
    if msg.text == "faq":
        bot.send_message(chat_id,
                     'Ответы на часто задаваемые вопросы вы можете найти [здесь](https://vk.com/@lvvlabel-chastye-voprosy)',
                     parse_mode='Markdown')    
    bot.register_next_step_handler(msg, on_click)

@bot.callback_query_handler(func=lambda callback: True)
def callback_message(callback):
    chat_id = callback.message.chat.id
    if callback.data == 'menu':
        keyboard = telebot.types.ReplyKeyboardMarkup()
        btn1 = telebot.types.KeyboardButton("Связаться с менеджером")
        keyboard.row(btn1)
        btn2 = telebot.types.KeyboardButton("Услуги")
        btn3 = telebot.types.KeyboardButton("Отзывы")
        keyboard.row(btn2, btn3)
        btn4 = telebot.types.KeyboardButton("О нас")
        btn5 = telebot.types.KeyboardButton("FAQ")
        keyboard.row(btn4, btn5)
        bot.send_message(chat_id,
                        'Вот что мы можем вам предложить',
                        reply_markup=keyboard)

def start_tg_bot():
    bot.polling(none_stop=True)