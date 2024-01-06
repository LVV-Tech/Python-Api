import vk_api
from vk_api.longpoll import VkLongPoll, VkEventType
from vk_api.keyboard import VkKeyboard, VkKeyboardColor
import os
from dotenv import load_dotenv, find_dotenv
import storage
import re
import json
import requests
from docx import Document
from datetime import datetime
import locale
locale.setlocale(locale.LC_ALL, "")


load_dotenv(find_dotenv())

token: str = os.getenv(key="VK_TOKEN")

vk_session: vk_api.VkApi = vk_api.VkApi(token=token)
vk = vk_session.get_api()
session_api = vk_session.get_api()
longpoll = VkLongPoll(vk_session)

services = ["дистрибьюция", "сведение и мастеринг", "создание дизайна", "аренда бита", "создание трека"]
servicesDescriptions = ["Дистрибуция - публикация Вашего трека/Альбома на все цифровые площадки.\nКаждый релиз отправляется на модерацию ПРОМО всех стриминговых сервисов.\nСсылка на портфолио - (позже)\n\n⏳ Сроки: от 3 до 5 дней.",
                        "Сведение и Мастеринг - выполнение комплекса работ по постобработке музыкального произведения.\nОчистка и подготовка звуковых дорожек, Обработка вокала, Питч-коррекция. Добавление различных эффектов и особенностей, Расширение стереобазы, работа с панорамой, Эквализация и компрессия,  и многое другое.\nВсе Ваши пожелания и требования будут учтены в процессе работы. Вам доступно Три бесплатных правки итоговой работы.\nСсылка на портфолио - (позже)\n\n⏳ Сроки: от 3 до 5 дней.",
                        "Создание дизайна - Создание обложки для Релиза/Альбома, Ретушь и обработка фото, создание дизайна для ПРОМО.\nВсе Ваши пожелания и требования будут учтены в процессе работы. Вам доступно Три бесплатных правки итоговой работы.\nСсылка на портфолио - https://vk.com/album-221967488_295470314\n\n⏳ Сроки: от 3 до 5 дней.",
                        "Аренда бита - любая категория аренды бита, от MP3 до EXCLUSIVE. Все биты представленные в плейлистах доступны для покупки.\nПлейлисты разбиты по разным Авторам и  жанрам, выбирайте подходящий Вам звук.\n\nСсылка на плейлист 1 - (позже)\n\n⏳ Сроки: 1 день.",
                        "Создание трека - услуга в полном объеме временно недоступна.\n\n⏳ Сроки: по договорённости"]
servicesPrices = [500, 2500, 1500, 1500, 5000]

def sender(id: int, text: str, keyboard: VkKeyboard = None):
    post = {"user_id": id, "message": text, "random_id": 0}
    if keyboard is not None:
        post["keyboard"] = keyboard.get_keyboard()
    vk_session.method("messages.send", post)

def get_last_msg(peer_id: int, msg_id: int):
    post = {"peer_id": peer_id, "conversation_message_ids": [msg_id+2]}
    return vk_session.method("messages.getByConversationMessageId", post)

def get_info(id: int):
    post = {"user_ids": [id]}
    return vk_session.method("users.get", post)

def start_vk_bot():
    print("Server in work")
    for event in longpoll.listen():
        if event.type == VkEventType.MESSAGE_NEW:
            
            if event.to_me:
                print(f'{event.text} from {event.user_id}')
                msg = event.text.lower()
                id = event.user_id
                
                if msg == "начать" or msg == "меню":
                    keyboard = VkKeyboard()
                    keyboard.add_button(
                        "Связаться с менеджером", VkKeyboardColor.POSITIVE
                    )
                    keyboard.add_line()
                    keyboard.add_button("Услуги", VkKeyboardColor.NEGATIVE)
                    keyboard.add_openlink_button("Отзывы", "https://vk.com/topic-221967488_49474380")
                    keyboard.add_line()
                    keyboard.add_openlink_button("О нас", "https://vk.com/@lvvlabel-kto-my")
                    keyboard.add_openlink_button("FAQ", "https://vk.com/@lvvlabel-chastye-voprosy")
                    sender(id, "Вот что мы можем вам предложить", keyboard)
                if msg == "регистрация":
                    keyboard = VkKeyboard()
                    keyboard.add_button("меню",VkKeyboardColor.SECONDARY)
                    sender(id, "Введите свои паспортные данные в формате: серия номер",keyboard)
                    creds = get_last_msg(event.peer_id, event.message_id)
                    while creds['count'] == 0:
                        creds = get_last_msg(event.peer_id, event.message_id)
                    print(creds)
                    match = re.search(r'^[0-9][0-9][0-9][0-9] [0-9][0-9][0-9][0-9][0-9][0-9]$', creds['items'][0]['text'])
                    if not match:
                        sender(id, "Ошибка, проверьте правильность введенных данных",keyboard)
                        continue
                    else:
                        #sender(id, "окок",keyboard)
                        passp = creds['items'][0]['text']
                    
                    sender(id, "Введите свой номер телефона в формате: +79999999999",keyboard)
                    creds = get_last_msg(event.peer_id, event.message_id+2)
                    while creds['count'] == 0:
                        creds = get_last_msg(event.peer_id, event.message_id+2)
                    print(creds)
                    match = re.search(r'^\+7[0-9][0-9][0-9][0-9][0-9][0-9][0-9][0-9][0-9][0-9]$', creds['items'][0]['text'])
                    if not match:
                        sender(id, "Ошибка, проверьте правильность введенных данных",keyboard)
                        continue
                    else:
                        ##sender(id, "окок",keyboard)
                        phone = creds['items'][0]['text']
                    
                    sender(id, "Введите своё имя фамилию отчество:",keyboard)
                    creds = get_last_msg(event.peer_id, event.message_id+4)
                    while creds['count'] == 0:
                        creds = get_last_msg(event.peer_id, event.message_id+4)
                    match = re.search(r'^[^\W\d_]+\s[^\W\d_]+?(\s[^\W\d_]+)$', creds['items'][0]['text'])
                    if not match:
                        sender(id, "Ошибка, проверьте правильность введенных данных",keyboard)
                        print(match, creds)
                        continue
                    else:
                        sender(id, "Вы зарегистрированы",keyboard)
                        name = creds['items'][0]['text']
                    
                    storage.create_user(
                        role=0,
                        phone=phone,
                        full_name=name,
                        passport=passp,
                        vk_id=id
                    )
                    
                        
                if msg == "связаться с менеджером":
                    keyboard = VkKeyboard()
                    keyboard.add_button("Творческий/Примеры работ, оценка, услуги", VkKeyboardColor.SECONDARY)
                    keyboard.add_line()
                    keyboard.add_button("Юридический вопрос/Оформление документов", VkKeyboardColor.SECONDARY)
                    keyboard.add_line()
                    keyboard.add_button("Маркетинг/Реклама, продвижение", VkKeyboardColor.SECONDARY)
                    keyboard.add_line()
                    keyboard.add_button("Технический вопрос/Не работает сайт, бот", VkKeyboardColor.SECONDARY)
                    keyboard.add_line()
                    keyboard.add_button("Меню", VkKeyboardColor.NEGATIVE)
                    sender(id, "Мы прочитаем ваш вопрос и ответим на него как можно скорее. Для удобства выберите тему вопроса.", keyboard)
                if msg == "творческий/примеры работ, оценка, услуги":
                    keyboard = VkKeyboard(inline=True)
                    keyboard.add_button("Меню", VkKeyboardColor.NEGATIVE)
                    sender(id, "Чем я могу вам помочь? Опишите свою проблему следующим сообщением", keyboard)
                if msg == "юридический вопрос/оформление документов":
                    keyboard = VkKeyboard(inline=True)
                    keyboard.add_button("Меню", VkKeyboardColor.NEGATIVE)
                    sender(id, "Чем я могу вам помочь? Опишите свою проблему следующим сообщением", keyboard)
                if msg == "маркетинг/реклама, продвижение":
                    keyboard = VkKeyboard(inline=True)
                    keyboard.add_button("Меню", VkKeyboardColor.NEGATIVE)
                    sender(id, "Чем я могу вам помочь? Опишите свою проблему следующим сообщением", keyboard)
                if msg == "технический вопрос/не работает сайт, бот":
                    keyboard = VkKeyboard(inline=True)
                    keyboard.add_button("Меню", VkKeyboardColor.NEGATIVE)
                    sender(id, "Чем я могу вам помочь? Опишите свою проблему следующим сообщением", keyboard)
                if msg == "услуги" or msg == "⏪назад⏪":
                    keyboard = VkKeyboard()
                    keyboard.add_button("Дистрибьюция", VkKeyboardColor.SECONDARY)
                    keyboard.add_line()
                    keyboard.add_button("Сведение и Мастеринг", VkKeyboardColor.SECONDARY)
                    keyboard.add_line()
                    keyboard.add_button("Создание Дизайна", VkKeyboardColor.SECONDARY)
                    keyboard.add_line()
                    keyboard.add_button("Аренда Бита", VkKeyboardColor.SECONDARY)
                    keyboard.add_line()
                    keyboard.add_button("Создание трека", VkKeyboardColor.SECONDARY)
                    keyboard.add_line()
                    keyboard.add_button("Узнать про пакеты", VkKeyboardColor.POSITIVE)
                    keyboard.add_line()
                    keyboard.add_button("Меню", VkKeyboardColor.NEGATIVE)
                    sender(id, "Выберите услугу", keyboard)
                for i in range(len(services)):
                    if msg == services[i]:
                        prevService = i
                        keyboard = VkKeyboard()
                        keyboard.add_button("✅Подтвердить заказ✅", VkKeyboardColor.POSITIVE)
                        keyboard.add_line()
                        keyboard.add_button("⏪Назад⏪")
                        sender(id, f"{servicesDescriptions[i]}\n💸 Цена: от {servicesPrices[i]} ₽.", keyboard)
                        break
                if msg == "узнать про пакеты":
                    keyboard = VkKeyboard()
                    keyboard.add_button("Меню", VkKeyboardColor.NEGATIVE)
                    sender(id, "Ваш Запрос в работе, в ближайшее время с вами свяжется менеджер", keyboard)
                if msg == "✅подтвердить заказ✅":
                    keyboard = VkKeyboard()
                    keyboard.add_button("Меню", VkKeyboardColor.NEGATIVE)
                    print("ВЫ выбрали услугу", services[prevService])
                    if storage.check_user(vk_id=id):
                        pass
                    else:
                        keyboard.add_button(
                            "Регистрация", VkKeyboardColor.POSITIVE
                        )
                        sender(id, "Вам необходимо зарегестрироваться", keyboard)
                        continue
                    
                    sender(id, "Заполните документы и отправьте их модератору: @ghostikgh", keyboard)
                    # Открываем документ
                    doc = Document("sogl.docx")
                    user = storage.get_user_vk_id(id)
                    # Получаем все параграфы документа
                    paras = doc.paragraphs
                    
                    today = datetime.now()
                    

                    # Проходим по всем параграфам и заменяем необходимые поля
                    for para in paras:
                        if "{{name}}" in para.text:
                            para.text = para.text.replace("{{name}}", user[0][2].title())
                    for para in paras:
                        if "{{passport}}" in para.text:
                            para.text = para.text.replace("{{passport}}", user[0][3])
                    for para in paras:
                        if "{{day}}" in para.text:
                            para.text = para.text.replace("{{day}}", datetime.strftime(today, '%d'))
                    for para in paras:
                        if "{{month}}" in para.text:
                            para.text = para.text.replace("{{month}}", datetime.strftime(today, ' %B'))
                    for para in paras:
                        if "{{year}}" in para.text:
                            para.text = para.text.replace("{{year}}", datetime.strftime(today, '%Y'))

                    # Сохраняем изменения
                    doc.save("output.docx")
                    result = json.loads(requests.post(vk.docs.getMessagesUploadServer(type='doc', peer_id=event.peer_id)['upload_url'],
                                                  files={'file': open('output.docx', 'rb')}).text)
                    jsonAnswer = vk.docs.save(file=result['file'], title=f'Согласие_обработку_{user[0][2].title()}', tags=[])

                    vk.messages.send(
                        peer_id=event.peer_id,
                        random_id=0,
                        attachment=f"doc{jsonAnswer['doc']['owner_id']}_{jsonAnswer['doc']['id']}"
                    )
                    
                    doc = Document("dog.docx")
                    user = storage.get_user_vk_id(id)
                    # Получаем все параграфы документа
                    paras = doc.paragraphs
                    
                    today = datetime.now()
                    

                    # Проходим по всем параграфам и заменяем необходимые поля
                    for para in paras:
                        if "{{name}}" in para.text:
                            para.text = para.text.replace("{{name}}", user[0][2].title())
                    for para in paras:
                        if "{{passport}}" in para.text:
                            para.text = para.text.replace("{{passport}}", user[0][3])
                    for para in paras:
                        if "{{day}}" in para.text:
                            para.text = para.text.replace("{{day}}", datetime.strftime(today, '%d'))
                    for para in paras:
                        if "{{month}}" in para.text:
                            para.text = para.text.replace("{{month}}", datetime.strftime(today, ' %B'))
                    for para in paras:
                        if "{{year}}" in para.text:
                            para.text = para.text.replace("{{year}}", datetime.strftime(today, '%Y'))
                    for para in paras:
                        if "{{services}}" in para.text:
                            para.text = para.text.replace("{{services}}", services[prevService])
                    for para in paras:
                        if "{{price}}" in para.text:
                            para.text = para.text.replace("{{price}}", str(servicesPrices[prevService]))

                    # Сохраняем изменения
                    doc.save("output.docx")
                    result = json.loads(requests.post(vk.docs.getMessagesUploadServer(type='doc', peer_id=event.peer_id)['upload_url'],
                                                  files={'file': open('output.docx', 'rb')}).text)
                    jsonAnswer = vk.docs.save(file=result['file'], title=f'Договор_{user[0][2].title()}', tags=[])

                    vk.messages.send(
                        peer_id=event.peer_id,
                        random_id=0,
                        attachment=f"doc{jsonAnswer['doc']['owner_id']}_{jsonAnswer['doc']['id']}"
                    )